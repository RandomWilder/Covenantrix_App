"""
Optimized Document Processor - Enhanced document processing with expanded format support
Handles multiple document formats with OCR support and legal document optimization
"""

import logging
import fitz  # pymupdf
import re
import magic
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import hashlib
import json

# Optional imports with fallback
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX support disabled.")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("openpyxl not available. Excel support disabled.")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("PIL/pytesseract not available. OCR support disabled.")


class DocumentProcessor:
    """
    Enhanced document processor with expanded format support and legal document optimization
    """
    
    def __init__(self, max_chunk_size: int = 4000, chunk_overlap: int = 200, enable_ocr: bool = False):
        """
        Initialize document processor with enhanced capabilities
        
        Args:
            max_chunk_size: Maximum size for document chunks
            chunk_overlap: Overlap between chunks for context preservation
            enable_ocr: Enable OCR for scanned documents and images
        """
        self.logger = logging.getLogger(__name__)
        
        # Core configuration
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        
        # Supported formats with capabilities
        self.supported_formats = {
            '.pdf': {'handler': 'extract_text_from_pdf', 'ocr_capable': True},
            '.txt': {'handler': 'extract_text_from_txt', 'ocr_capable': False},
            '.docx': {'handler': 'extract_text_from_docx', 'ocr_capable': False, 'available': DOCX_AVAILABLE},
            '.doc': {'handler': 'extract_text_from_docx', 'ocr_capable': False, 'available': DOCX_AVAILABLE},
            '.xlsx': {'handler': 'extract_text_from_excel', 'ocr_capable': False, 'available': EXCEL_AVAILABLE},
            '.xls': {'handler': 'extract_text_from_excel', 'ocr_capable': False, 'available': EXCEL_AVAILABLE},
            '.png': {'handler': 'extract_text_from_image', 'ocr_capable': True, 'available': OCR_AVAILABLE},
            '.jpg': {'handler': 'extract_text_from_image', 'ocr_capable': True, 'available': OCR_AVAILABLE},
            '.jpeg': {'handler': 'extract_text_from_image', 'ocr_capable': True, 'available': OCR_AVAILABLE},
            '.tiff': {'handler': 'extract_text_from_image', 'ocr_capable': True, 'available': OCR_AVAILABLE},
            '.tif': {'handler': 'extract_text_from_image', 'ocr_capable': True, 'available': OCR_AVAILABLE}
        }
        
        # Legal document patterns for enhanced processing
        self.legal_patterns = {
            'contract_parties': re.compile(r'(?:PARTIES?|BETWEEN|LANDLORD|TENANT|BUYER|SELLER|CLIENT|CONTRACTOR):\s*([^\n]+)', re.IGNORECASE),
            'contract_dates': re.compile(r'(?:DATE|DATED|EFFECTIVE|EXECUTION|COMMENCEMENT):\s*([^\n]+)', re.IGNORECASE),
            'monetary_amounts': re.compile(r'\$[\d,]+\.?\d*|\d+\s*(?:dollars?|USD)', re.IGNORECASE),
            'legal_sections': re.compile(r'(?:ARTICLE|SECTION|CLAUSE)\s+[\d\w]+[:\.]', re.IGNORECASE),
            'addresses': re.compile(r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd)[^\n]*', re.IGNORECASE)
        }
        
        self.logger.info(f"Document processor initialized - OCR: {self.enable_ocr}")
        self._log_available_formats()
    
    def process_document_file(self, file_path: str, document_type: str = "contract") -> Dict[str, Any]:
        """
        Enhanced document processing with metadata extraction and legal document optimization
        
        Args:
            file_path: Path to the document file
            document_type: Type of document (contract, legal, general)
            
        Returns:
            Dict containing processing results with enhanced metadata
        """
        try:
            path = Path(file_path)
            
            if not path.exists():
                return {"success": False, "error": "File not found"}
            
            # Validate file format
            file_extension = path.suffix.lower()
            format_info = self.supported_formats.get(file_extension)
            
            if not format_info:
                return {
                    "success": False,
                    "error": f"Unsupported format: {file_extension}",
                    "supported_formats": list(self.supported_formats.keys())
                }
            
            # Check if format is available
            if format_info.get('available', True) is False:
                return {
                    "success": False,
                    "error": f"Format {file_extension} requires additional dependencies",
                    "install_hint": self._get_install_hint(file_extension)
                }
            
            # Extract file metadata
            file_metadata = self._extract_file_metadata(path)
            
            # Extract text using appropriate handler
            handler_name = format_info['handler']
            handler = getattr(self, handler_name)
            text = handler(str(path))
            
            if not text.strip():
                return {
                    "success": False,
                    "error": "No text extracted from document",
                    "file_metadata": file_metadata
                }
            
            # Enhanced metadata extraction
            document_metadata = self._extract_document_metadata(text, document_type)
            
            # Apply intelligent chunking
            chunks = self.chunk_text_intelligently(text, document_type)
            
            # Calculate document hash for deduplication
            document_hash = self._calculate_document_hash(text)
            
            result = {
                "success": True,
                "text": text,
                "chunks": chunks,
                "filename": path.name,
                "format": file_extension,
                "document_type": document_type,
                "document_hash": document_hash,
                "file_metadata": file_metadata,
                "document_metadata": document_metadata,
                "processing_stats": {
                    "char_count": len(text),
                    "word_count": len(text.split()),
                    "chunk_count": len(chunks),
                    "chunking_applied": len(chunks) > 1,
                    "processed_at": datetime.now().isoformat()
                }
            }
            
            self.logger.info(f"✅ Document processed: {path.name} ({len(chunks)} chunks)")
            return result
            
        except Exception as e:
            self.logger.error(f"❌ Document processing failed for {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": Path(file_path).name,
                "traceback": str(e)
            }
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Enhanced PDF extraction with OCR fallback for scanned documents"""
        try:
            doc = fitz.open(file_path)
            text = ""
            pages_with_text = 0
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                # If page has little text and OCR is enabled, try OCR
                if len(page_text.strip()) < 50 and self.enable_ocr:
                    try:
                        # Convert page to image and OCR
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        
                        # Save temporarily and OCR
                        import tempfile
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_file:
                            tmp_file.write(img_data)
                            tmp_file.flush()
                            
                            ocr_text = pytesseract.image_to_string(Image.open(tmp_file.name))
                            if len(ocr_text.strip()) > len(page_text.strip()):
                                page_text = ocr_text
                                self.logger.debug(f"OCR applied to page {page_num + 1}")
                            
                            Path(tmp_file.name).unlink()  # Cleanup
                            
                    except Exception as ocr_error:
                        self.logger.warning(f"OCR failed for page {page_num + 1}: {ocr_error}")
                
                if page_text.strip():
                    pages_with_text += 1
                
                text += f"[Page {page_num + 1}]\n{page_text}\n\n"
            
            doc.close()
            
            self.logger.debug(f"PDF processed: {pages_with_text} pages with text")
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"PDF extraction failed for {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Enhanced TXT extraction with encoding detection"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        self.logger.debug(f"TXT file read with encoding: {encoding}")
                        return content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                content = file.read().decode('utf-8', errors='ignore')
                self.logger.warning(f"TXT file read with error handling - some characters may be lost")
                return content.strip()
                
        except Exception as e:
            self.logger.error(f"TXT extraction failed for {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX/DOC files with table support"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX support: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files with sheet support"""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel support: pip install openpyxl")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"[SHEET: {sheet_name}]"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None and str(cell).strip():
                            row_text.append(str(cell).strip())
                    
                    if row_text:
                        sheet_text.append(" | ".join(row_text))
                
                if len(sheet_text) > 1:  # More than just the sheet header
                    text_parts.append("\n".join(sheet_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Excel extraction failed for {file_path}: {e}")
            return ""
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        if not OCR_AVAILABLE:
            raise ImportError("PIL and pytesseract are required for image OCR: pip install Pillow pytesseract")
        
        try:
            image = Image.open(file_path)
            
            # Optimize image for OCR
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Apply OCR with custom config for better results
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Image OCR failed for {file_path}: {e}")
            return ""
    
    def chunk_text_intelligently(self, text: str, document_type: str = "general") -> List[str]:
        """
        Enhanced intelligent chunking with legal document optimization
        
        Args:
            text: The text to chunk
            document_type: Type of document for specialized chunking
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.max_chunk_size:
            return [text]
        
        chunks = []
        
        # Legal document specific chunking
        if document_type in ["contract", "legal"]:
            chunks = self._chunk_legal_document(text)
        else:
            chunks = self._chunk_general_document(text)
        
        # Apply overlap strategy
        if len(chunks) > 1:
            chunks = self._add_chunk_overlap(chunks)
        
        self.logger.info(f"📄 Text chunked into {len(chunks)} intelligent chunks ({document_type})")
        return chunks
    
    def _chunk_legal_document(self, text: str) -> List[str]:
        """Specialized chunking for legal documents"""
        chunks = []
        
        # Try to split by legal sections first
        section_pattern = re.compile(r'(?:ARTICLE|SECTION|CLAUSE)\s+[\d\w]+[:\.].*?(?=(?:ARTICLE|SECTION|CLAUSE)\s+[\d\w]+[:.]|$)', 
                                   re.DOTALL | re.IGNORECASE)
        sections = section_pattern.findall(text)
        
        if sections and len(sections) > 1:
            # Use section-based chunking
            current_chunk = ""
            
            for section in sections:
                if len(current_chunk) + len(section) <= self.max_chunk_size:
                    current_chunk += section + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = section + "\n\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        else:
            # Fallback to paragraph-based chunking
            chunks = self._chunk_general_document(text)
        
        return chunks
    
    def _chunk_general_document(self, text: str) -> List[str]:
        """General document chunking strategy"""
        chunks = []
        
        # Split by double newlines (paragraphs)
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for paragraph in paragraphs:
            # If adding this paragraph would exceed chunk size
            if len(current_chunk) + len(paragraph) + 2 > self.max_chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    
                # If paragraph itself is too long, split by sentences
                if len(paragraph) > self.max_chunk_size:
                    sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                    temp_chunk = ""
                    
                    for sentence in sentences:
                        if len(temp_chunk) + len(sentence) + 1 > self.max_chunk_size:
                            if temp_chunk:
                                chunks.append(temp_chunk.strip())
                            temp_chunk = sentence
                        else:
                            temp_chunk += " " + sentence if temp_chunk else sentence
                    
                    current_chunk = temp_chunk
                else:
                    current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph
        
        # Add the final chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _add_chunk_overlap(self, chunks: List[str]) -> List[str]:
        """Add intelligent overlap between chunks"""
        overlapped_chunks = []
        
        for i, chunk in enumerate(chunks):
            if i == 0:
                overlapped_chunks.append(chunk)
            else:
                # Get overlap from previous chunk
                prev_chunk = chunks[i-1]
                overlap_text = prev_chunk[-self.chunk_overlap:] if len(prev_chunk) > self.chunk_overlap else prev_chunk
                
                # Find a good break point in overlap (sentence boundary)
                sentences = re.split(r'(?<=[.!?])\s+', overlap_text)
                if len(sentences) > 1:
                    # Use the last few complete sentences for overlap
                    overlap_text = '. '.join(sentences[-2:]) if len(sentences) > 2 else sentences[-1]
                
                overlapped_chunks.append(f"[CONTEXT FROM PREVIOUS SECTION]\n{overlap_text}\n\n{chunk}")
        
        return overlapped_chunks
    
    def _extract_file_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract file system metadata"""
        try:
            stat = path.stat()
            
            # Detect MIME type
            mime_type = "unknown"
            try:
                mime_type = magic.from_file(str(path), mime=True)
            except:
                pass
            
            return {
                "filename": path.name,
                "file_size": stat.st_size,
                "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "extension": path.suffix.lower(),
                "mime_type": mime_type
            }
        except Exception as e:
            return {"error": str(e)}
    
    def _extract_document_metadata(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract document content metadata"""
        metadata = {
            "document_type": document_type,
            "language": "en",  # Could be enhanced with language detection
            "extracted_entities": {}
        }
        
        # Legal document specific extraction
        if document_type in ["contract", "legal"]:
            for pattern_name, pattern in self.legal_patterns.items():
                matches = pattern.findall(text)
                if matches:
                    metadata["extracted_entities"][pattern_name] = matches[:5]  # Limit to 5 matches
        
        # General document analysis
        metadata.update({
            "paragraph_count": len([p for p in text.split('\n\n') if p.strip()]),
            "sentence_count": len(re.findall(r'[.!?]+', text)),
            "contains_tables": "[TABLE]" in text,
            "contains_monetary_amounts": bool(self.legal_patterns['monetary_amounts'].search(text))
        })
        
        return metadata
    
    def _calculate_document_hash(self, text: str) -> str:
        """Calculate hash for document deduplication"""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()[:16]
    
    def _get_install_hint(self, file_extension: str) -> str:
        """Get installation hint for missing dependencies"""
        hints = {
            '.docx': "pip install python-docx",
            '.doc': "pip install python-docx",
            '.xlsx': "pip install openpyxl",
            '.xls': "pip install openpyxl",
            '.png': "pip install google-cloud-vision (and configure Google Vision API key)",
            '.jpg': "pip install google-cloud-vision (and configure Google Vision API key)",
            '.jpeg': "pip install google-cloud-vision (and configure Google Vision API key)",
            '.tiff': "pip install google-cloud-vision (and configure Google Vision API key)",
            '.tif': "pip install google-cloud-vision (and configure Google Vision API key)"
        }
        return hints.get(file_extension, "Check documentation for required dependencies")
    
    def _log_available_formats(self):
        """Log available and unavailable formats"""
        available = []
        unavailable = []
        ocr_formats = []
        
        for ext, info in self.supported_formats.items():
            if info.get('available', True):
                available.append(ext)
                if info.get('ocr_capable', False) and self.enable_ocr:
                    ocr_formats.append(ext)
            else:
                unavailable.append(ext)
        
        self.logger.info(f"Available formats: {', '.join(available)}")
        if ocr_formats:
            self.logger.info(f"OCR-enabled formats: {', '.join(ocr_formats)}")
        if unavailable:
            self.logger.warning(f"Unavailable formats (missing deps): {', '.join(unavailable)}")
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported and available"""
        ext = Path(filename).suffix.lower()
        format_info = self.supported_formats.get(ext)
        return format_info is not None and format_info.get('available', True)
    
    def requires_ocr(self, filename: str) -> bool:
        """Check if file format requires OCR for processing"""
        ext = Path(filename).suffix.lower()
        format_info = self.supported_formats.get(ext)
        if not format_info:
            return False
        
        # Image formats always require OCR
        if ext in ['.png', '.jpg', '.jpeg', '.tiff', '.tif']:
            return True
        
        # PDFs might require OCR for scanned pages
        if ext == '.pdf':
            return True  # We'll detect this during processing
        
        return False
    
    def get_supported_formats(self) -> Dict[str, Any]:
        """Get comprehensive information about supported formats"""
        return {
            ext: {
                "available": info.get('available', True),
                "ocr_capable": info.get('ocr_capable', False),
                "ocr_enabled": info.get('ocr_capable', False) and self.enable_ocr,
                "requires_ocr": ext in ['.png', '.jpg', '.jpeg', '.tiff', '.tif'],
                "install_hint": self._get_install_hint(ext) if not info.get('available', True) else None
            }
            for ext, info in self.supported_formats.items()
        }